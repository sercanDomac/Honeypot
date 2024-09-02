from flask import Flask, send_file, request, abort
import os
import random
from faker import Faker
import pandas as pd
from fpdf import FPDF
from docx import Document
import logging
import time

app = Flask(__name__)

# Logger ayarları (saldırgan aktivitelerini kaydetmek için)
logging.basicConfig(filename='honeypot.log', level=logging.INFO, format='%(asctime)s - %(message)s')

# Faker nesnesi oluşturma ve Türkçe dilini kullanma
faker = Faker('tr_TR')

# Sahte Veeam yedekleme dosyaları üretimi
def generate_fake_veeam_backup_files():
    backup_files = []
    for _ in range(100):  # 100 adet Veeam yedekleme dosyası üretelim
        job_name = faker.word() + "_yedekleme_görevi"
        file_name = f"{job_name}_{random.randint(1, 100)}.vbk"
        content = f"Yedekleme Görevi: {job_name}\nOluşturulma Tarihi: {faker.date_this_year()}\nDosya Boyutu: {random.randint(50, 500)} GB\nYedeklenen Sunucular: {faker.company_suffix()} Sunucusu"
        backup_files.append({"dosya_adı": file_name, "içerik": content})
        
        # Dosyayı kaydet
        file_path = os.path.abspath(os.path.join("fake_data", file_name))
        with open(file_path, 'w') as file:
            file.write(content)
            
    return backup_files

# Şirket hakkında PDF dosyası üretimi
def generate_fake_pdf_report():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Şirket Raporu", ln=True, align='C')
    pdf.cell(200, 10, txt=f"Şirket Adı: {faker.company()}", ln=True)
    pdf.cell(200, 10, txt=f"Adres: {faker.address()}", ln=True)
    pdf.cell(200, 10, txt=f"Kuruluş Tarihi: {faker.date()}", ln=True)
    pdf.cell(200, 10, txt=f"Faaliyet Alanı: {faker.bs()}", ln=True)
    
    for i in range(1, 10):  # Daha fazla rapor bölümü ekleyelim
        pdf.multi_cell(200, 10, txt=f"Rapor Bölümü {i}: {faker.text(max_nb_chars=500)}", align='L')
    
    pdf_output = os.path.abspath(os.path.join("fake_data", "sirket_raporu.pdf"))
    pdf.output(pdf_output)
    return pdf_output

# Şirket hakkında Excel dosyası üretimi
def generate_fake_excel_data():
    data = {
        "Çalışan Adı": [faker.name() for _ in range(200)],  # 100 yerine 200 çalışan
        "Departman": [faker.job() for _ in range(200)],
        "Maaş": [f"{faker.random_int(min=30000, max=100000):,} TL" for _ in range(200)],
        "Giriş Tarihi": [faker.date_this_decade() for _ in range(200)],
    }
    df = pd.DataFrame(data)
    excel_output = os.path.abspath(os.path.join("fake_data", "calisan_verileri.xlsx"))
    df.to_excel(excel_output, index=False)
    return excel_output

# Şirket hakkında Word belgesi üretimi
def generate_fake_word_document():
    doc = Document()
    doc.add_heading('Şirket Bilgileri', 0)
    
    doc.add_paragraph(f"Şirket Adı: {faker.company()}")
    doc.add_paragraph(f"Adres: {faker.address()}")
    doc.add_paragraph(f"Kuruluş Tarihi: {faker.date()}")
    doc.add_paragraph(f"Faaliyet Alanı: {faker.catch_phrase()}")
    
    doc.add_heading('Önemli Notlar', level=1)
    for _ in range(10):  # Daha fazla önemli not ekleyelim
        doc.add_paragraph(faker.text(max_nb_chars=500))
    
    word_output = os.path.abspath(os.path.join("fake_data", "sirket_bilgileri.docx"))
    doc.save(word_output)
    return word_output

# Sahte dosyaları kaydetme dizini
os.makedirs("fake_data", exist_ok=True)

# Ana sayfa rotası
@app.route('/')
def index():
    logging.info(f"Erişim kaydedildi: {request.remote_addr} tarafından")
    return '''
    <h1>Honeypot Sahte Sunucusu</h1>
    <p>Aşağıdaki sahte verilere erişebilirsiniz:</p>
    <ul>
        <li><a href="/veeam">Veeam Yedekleme Dosyaları</a></li>
        <li><a href="/pdf">Şirket Raporu (PDF)</a></li>
        <li><a href="/excel">Çalışan Verileri (Excel)</a></li>
        <li><a href="/word">Şirket Bilgileri (Word)</a></li>
    </ul>
    '''

# Veeam yedekleme dosyaları rotası
@app.route('/veeam')
def veeam():
    data = generate_fake_veeam_backup_files()
    logging.info(f"Veeam yedekleme dosyalarına erişim: {request.remote_addr} tarafından")

    # Rastgele olarak sahte hata mesajı gönderme
    if random.choice([True, False]):
        logging.warning(f"Sahte hata mesajı gösterildi: {request.remote_addr}")
        abort(500, description="Sunucu Hatası: Yedekleme dosyasına erişim sırasında bir sorun oluştu.")
    
    # Rastgele dosya erişimini yavaşlatma
    if random.choice([True, False]):
        logging.info(f"Erişim yavaşlatıldı: {request.remote_addr}")
        time.sleep(random.uniform(2, 5))  # 2 ile 5 saniye arasında bekleme
    
    html_content = "<ul>"
    for file in data:
        html_content += f"<li>{file['dosya_adı']}: {file['içerik'][:50]}...</li>"
    html_content += "</ul>"
    return html_content

# PDF raporu rotası
@app.route('/pdf')
def pdf_report():
    file_path = generate_fake_pdf_report()
    logging.info(f"PDF raporuna erişim: {request.remote_addr} tarafından")

    # Rastgele olarak sahte hata mesajı gönderme
    if random.choice([True, False]):
        logging.warning(f"Sahte hata mesajı gösterildi: {request.remote_addr}")
        abort(500, description="Sunucu Hatası: PDF dosyasına erişim sırasında bir sorun oluştu.")
    
    # Rastgele dosya erişimini yavaşlatma
    if random.choice([True, False]):
        logging.info(f"Erişim yavaşlatıldı: {request.remote_addr}")
        time.sleep(random.uniform(2, 5))  # 2 ile 5 saniye arasında bekleme
    
    return send_file(file_path, as_attachment=True)

# Excel dosyası rotası
@app.route('/excel')
def excel_data():
    file_path = generate_fake_excel_data()
    logging.info(f"Excel dosyasına erişim: {request.remote_addr} tarafından")

    # Rastgele olarak sahte hata mesajı gönderme
    if random.choice([True, False]):
        logging.warning(f"Sahte hata mesajı gösterildi: {request.remote_addr}")
        abort(500, description="Sunucu Hatası: Excel dosyasına erişim sırasında bir sorun oluştu.")
    
    # Rastgele dosya erişimini yavaşlatma
    if random.choice([True, False]):
        logging.info(f"Erişim yavaşlatıldı: {request.remote_addr}")
        time.sleep(random.uniform(2, 5))  # 2 ile 5 saniye arasında bekleme
    
    return send_file(file_path, as_attachment=True)

# Word belgesi rotası
@app.route('/word')
def word_document():
    file_path = generate_fake_word_document()
    logging.info(f"Word belgesine erişim: {request.remote_addr} tarafından")

    # Rastgele olarak sahte hata mesajı gönderme
    if random.choice([True, False]):
        logging.warning(f"Sahte hata mesajı gösterildi: {request.remote_addr}")
        abort(500, description="Sunucu Hatası: Word dosyasına erişim sırasında bir sorun oluştu.")
    
    # Rastgele dosya erişimini yavaşlatma
    if random.choice([True, False]):
        logging.info(f"Erişim yavaşlatıldı: {request.remote_addr}")
        time.sleep(random.uniform(2, 5))  # 2 ile 5 saniye arasında bekleme
    
    return send_file(file_path, as_attachment=True)

# Uygulamayı çalıştırma
if __name__ == '__main__':
    app.run(debug=True)
